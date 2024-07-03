#!/usr/bin/env python3
"""将某一目录下（包括子文件夹）的所有Word文件中的表格数据汇总合并为一个xlsx文件

Example:
  合并Word中的表格为Excel ::

    doc_docx_dir = Path.cwd() / 'testdoc' # 处理 testdoc/ 目录下的所有word文件

    docfiles = gen_find('*.doc', doc_docx_dir)
    doc_to_docx(docfiles)  # 将doc文件转换为docx

    docxfiles = gen_find('*.docx', doc_docx_dir)
    df = process_docx(docxfiles, False)
    df.to_excel('汇总.xlsx', index=False, header=False)
"""

import platform
import subprocess
from pathlib import Path

import pandas as pd
from docx import Document


def gen_find(filepat, topdir):
    """Find all filenames in a directory tree that match a shell wildcard pattern

    Args:
        filepat: string, 通配符后缀名，如 ``'*.jpeg'``
        topdir: 递归查找的顶级目录

    Returns:
        generator(Path): 匹配到的文件Path生成器
    """
    p = Path(topdir)
    return p.glob('**/' + filepat)


def doc_to_docx(doc_files):
    """将一组doc格式文件批量转换为docx文件

    Args:
        doc_files: doc文件名列表、集合、元组或生成器

    Note:
        Windows需安装MS Office较新版；Linux需安装LibreOffice才能完成转换。

        暂不支持Mac系统。
    """
    for filename in doc_files:
        if platform.system() == 'Windows':
            from doc2docx import convert
            convert(filename)
        elif platform.system() == 'Linux':
            subprocess.call([
                'soffice', '--headless', '--convert-to', 'docx', filename,
                '--outdir', f'{filename.parents[0]}'
            ])
        else:
            print(
                f"error: doc_to_docx only support Win/Linux now.\nCan't support {platform.system()} now"
            )


def read_table_to_oneline(filename):
    """将一个docx文件中的所有table转换为单行数据

    Args:
        filename (Path): docx文件名

    Note:
        适用于表格操作或编程水平较好的用户。

        合并之后，还需要调整表格,将行中的 key列转换为列标题。
    """
    doc = Document(filename)
    data = [filename]

    # all_paras = doc.paragraphs
    # for para in all_paras:
    #     print(para.text)
    for table in doc.tables:
        # https://programmer.ink/think/python-reads-the-merged-cell-information-in-the-docx-table.html
        cells = table._cells
        for i, cell in enumerate(cells):
            # Skip if the cell does not appear for the first time in the table
            if cell in cells[:i]:
                continue
            else:
                data.append(cell.text)
    return data


def read_table_to_mulline(filename):
    """将一个docx文件中的所有table转换为多行数据列表

    Note:
        适用于普通用户。
    """
    doc = Document(filename)
    data = []

    # all_paras = doc.paragraphs
    # for para in all_paras:
    #     print(para.text)
    for table in doc.tables:
        for row in table.rows:
            rowdata = []
            row_cells = []
            for cell in row.cells:
                # 因为row.cells will return a "merged" cell multiple times, once
                # for each cell that is merged into iterate.

                # 一些word文件中多合并单元格。为应对其非计算机数据逻辑的使用方式，
                # 我们将每个表格中的内容简单输出为一行数据

                if cell in row_cells:
                    continue
                row_cells.append(cell)
                rowdata.append(cell.text)

            data.append(rowdata)
    return data


def process_docx(files, is_oneline=True):
    """将一系列docx文件中的所有表格转换到xlsx文件

    Note:
        适用于表格操作水平一般的用户。

        合并之后，还需要调整表格,将行中的key列转换为列标题。

    Args:
        files: docx文件列表或生成器
        is_oneline: True则将一个docx文件中的表格数据转换为xlsx中一行数据；False则为多行。

    Returns:
        DataFrame: Pandas的DataFrame格式。
    """
    datas = []
    datas.append([])
    print('{files}')
    for i in files:
        if is_oneline:
            datas.append(read_table_to_oneline(i))
        else:
            datas += read_table_to_mulline(i)

    df = pd.DataFrame.from_records(datas)
    return df
